"""Phase A1, Node 1: load a resume or JD and chunk it for extraction.

Deliberately dependency-light: `.txt`/`.md` need nothing, `.docx` needs
python-docx and only imports it when a .docx is actually opened, so the rest of
the pipeline stays importable on a machine without it.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

# ~500 tokens; English prose runs roughly 0.75 words per token, so ~375 words.
DEFAULT_CHUNK_WORDS = 375
DEFAULT_OVERLAP_WORDS = 40

SUPPORTED_SUFFIXES = {".txt", ".md", ".docx"}

_WHITESPACE = re.compile(r"[ \t\r\f\v]+")
_BLANKLINES = re.compile(r"\n{3,}")


@dataclass(frozen=True)
class Chunk:
    index: int
    text: str
    source_id: str


@dataclass
class Document:
    source_id: str
    path: Path
    text: str
    doc_type: str = "unknown"
    chunks: list[Chunk] = field(default_factory=list)


def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _read_docx(path: Path) -> str:
    try:
        import docx  # type: ignore
    except ImportError as exc:  # pragma: no cover - environment-dependent
        raise ImportError(
            "Reading .docx requires python-docx. Install it with `pip install python-docx`."
        ) from exc

    document = docx.Document(str(path))
    parts: list[str] = [p.text for p in document.paragraphs]
    # Resumes frequently hide the entire skills section inside a table.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def normalize_text(raw: str) -> str:
    """Collapse intra-line whitespace and runs of blank lines; keep line structure.

    Line breaks are load-bearing in resumes (they separate bullets), so they are
    preserved rather than flattened into a single paragraph.
    """
    lines = [_WHITESPACE.sub(" ", line).strip() for line in raw.splitlines()]
    return _BLANKLINES.sub("\n\n", "\n".join(lines)).strip()


def chunk_text(
    text: str,
    source_id: str,
    chunk_words: int = DEFAULT_CHUNK_WORDS,
    overlap_words: int = DEFAULT_OVERLAP_WORDS,
) -> list[Chunk]:
    """Split into overlapping word windows.

    Overlap exists so a skill named at a chunk boundary is not severed from the
    sentence that evidences it.
    """
    if chunk_words <= 0:
        raise ValueError("chunk_words must be positive")
    if not 0 <= overlap_words < chunk_words:
        raise ValueError("overlap_words must be >= 0 and < chunk_words")

    words = text.split()
    if not words:
        return []

    stride = chunk_words - overlap_words
    chunks: list[Chunk] = []
    for start in range(0, len(words), stride):
        window = words[start : start + chunk_words]
        if not window:
            break
        chunks.append(
            Chunk(index=len(chunks), text=" ".join(window), source_id=source_id)
        )
        if start + chunk_words >= len(words):
            break
    return chunks


def read_document(
    path: str | Path,
    doc_type: str = "unknown",
    chunk_words: int = DEFAULT_CHUNK_WORDS,
    overlap_words: int = DEFAULT_OVERLAP_WORDS,
) -> Document:
    """Load one file and return it normalized and chunked."""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(path)
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(
            f"Unsupported file type {suffix!r}; supported: {sorted(SUPPORTED_SUFFIXES)}"
        )

    raw = _read_docx(path) if suffix == ".docx" else _read_txt(path)
    text = normalize_text(raw)
    source_id = path.stem

    return Document(
        source_id=source_id,
        path=path,
        text=text,
        doc_type=doc_type,
        chunks=chunk_text(text, source_id, chunk_words, overlap_words),
    )


def read_directory(
    directory: str | Path, doc_type: str = "unknown", **kwargs: object
) -> list[Document]:
    """Read every supported file in a directory, sorted for reproducibility."""
    directory = Path(directory)
    paths = sorted(
        p for p in directory.iterdir() if p.suffix.lower() in SUPPORTED_SUFFIXES
    )
    return [read_document(p, doc_type=doc_type, **kwargs) for p in paths]  # type: ignore[arg-type]
